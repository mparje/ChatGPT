import openai
import streamlit as st
import pdfplumber
import io
from docx import Document

def extract_text_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# generate a response
def generate_response(prompt, text):
    st.session_state['messages'].append({"role": "user", "content": prompt})

    completion = openai.ChatCompletion.create(
        model=model,
        messages=st.session_state['messages']
    )
    response = completion.choices[0].message.content
    st.session_state['messages'].append({"role": "assistant", "content": response})
    
    completion = openai.chatCompletion.create(
        engine="gpt-4",
        prompt=prompt,
        max_tokens=250,
        n=1,
        stop=None,
        temperature=0.5,
        top_p=1,
        frequency_penalty=0.3,
        presence_penalty=0.3,
        echo=False
    )
    
    response = completion.choices[0].text.strip()
    return response

def handle_file_upload():
    if "answer" not in st.session_state:
        st.session_state["answer"] = {}

    st.title("Ask my Documents")
    
    left_column = st.sidebar
    left_column.title("Settings and Instructions")
    api_key = left_column.text_input("Enter your OpenAI API key:", type="password")
    if api_key == "":
        st.warning("Please enter a valid OpenAI API key.")
        return
    
    openai.api_key = api_key

    left_column.header("Instructions")
    left_column.write("This tool uses the GPT-4 model from OpenAI to answer questions based on uploaded documents. The documents can be in PDF or Docx format and should not exceed 4000 words or approximately 8 pages. The documents will be deleted from the server after 10 minutes of inactivity.")
    left_column.write("Note: You must have an API key corresponding to an account authorized to use GPT-4.")
    
    file = st.file_uploader("Upload a PDF or Docx file (max 2 MB)", type=["pdf", "docx"], accept_multiple_files=False)
    if file is not None:
        if file.size <= 2 * 1024 * 1024:
            if file.type == "application/pdf":
                text = extract_text_from_pdf(io.BytesIO(file.read()))
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = extract_text_from_docx(io.BytesIO(file.read()))
            
            if len(text.split()) > 4000:
                st.warning("The document exceeds the maximum word count (4000 words). Please upload a smaller document.")
                return

            custom_question = st.text_input("Enter your question:")
            if st.button("Submit question"):
                answer = generate_response(custom_question, text)
                st.write(answer)
        else:
            st.warning("The file size exceeds the maximum limit. Please upload a smaller file.")

if __name__ == "__main__":
    handle_file_upload()
import openai
import streamlit as st
import pdfplumber
import io
from docx import Document

def extract_text_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# generate a response
def generate_response(prompt, text):
    st.session_state['messages'].append({"role": "user", "content": prompt})

    completion = openai.ChatCompletion.create(
        model=model,
        messages=st.session_state['messages']
    )
    response = completion.choices[0].message.content
    st.session_state['messages'].append({"role": "assistant", "content": response})
    
    completion = openai.chatCompletion.create(
        engine="gpt-4",
        prompt=prompt,
        max_tokens=250,
        n=1,
        stop=None,
        temperature=0.5,
        top_p=1,
        frequency_penalty=0.3,
        presence_penalty=0.3,
        echo=False
    )
    
    response = completion.choices[0].text.strip()
    return response

def handle_file_upload():
    if "answer" not in st.session_state:
        st.session_state["answer"] = {}

    st.title("Ask my Documents")
    
    left_column = st.sidebar
    left_column.title("Settings and Instructions")
    api_key = left_column.text_input("Enter your OpenAI API key:", type="password")
    if api_key == "":
        st.warning("Please enter a valid OpenAI API key.")
        return
    
    openai.api_key = api_key

    left_column.header("Instructions")
    left_column.write("This tool uses the GPT-4 model from OpenAI to answer questions based on uploaded documents. The documents can be in PDF or Docx format and should not exceed 4000 words or approximately 8 pages. The documents will be deleted from the server after 10 minutes of inactivity.")
    left_column.write("Note: You must have an API key corresponding to an account authorized to use GPT-4.")
    
    file = st.file_uploader("Upload a PDF or Docx file (max 2 MB)", type=["pdf", "docx"], accept_multiple_files=False)
    if file is not None:
        if file.size <= 2 * 1024 * 1024:
            if file.type == "application/pdf":
                text = extract_text_from_pdf(io.BytesIO(file.read()))
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = extract_text_from_docx(io.BytesIO(file.read()))
            
            if len(text.split()) > 4000:
                st.warning("The document exceeds the maximum word count (4000 words). Please upload a smaller document.")
                return

            custom_question = st.text_input("Enter your question:")
            if st.button("Submit question"):
                answer = generate_response(custom_question, text)
                st.write(answer)
        else:
            st.warning("The file size exceeds the maximum limit. Please upload a smaller file.")

if __name__ == "__main__":
    handle_file_upload()
